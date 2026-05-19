"""
QuickCast 네이버 포스트용 DOCX — Step 4 (최종)
섹션 09 알람
섹션 10 키 입력 방식 선택
섹션 11 텔레그램 · 편의 기능
섹션 12 자주 묻는 질문
백 커버 (다운로드 / 한 줄 마무리 / 태그)
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from make_naverpost_step1 import (
    FONT_KR, FONT_EN, INK, INK_SOFT, ACCENT, RULE, GOLD,
    add_paragraph, add_page_break, style_run, set_cell_shading,
    make_cover, make_intro, setup_page,
)
from make_naverpost_step2 import (
    section_header, lead, body, h3, bullet, capture_placeholder,
    info_box, compare_table, link_bullet, add_hyperlink,
    section_01_compare, section_02_install, section_03_power, section_04_hpmp,
)
from make_naverpost_step3 import (
    section_05_slots, section_06_pk_potion,
    section_07_pet_whistle, section_08_recovery,
)


# ───────── 섹션 09 알람 ─────────
def section_09_alerts(doc):
    section_header(
        doc, "09", "알람",
        "시간 기반 알림 — 운영 점검, 결정전, 보스 출현 시간.",
    )
    lead(
        doc,
        "사냥 외에 매크로가 도와줄 수 있는 작은 일 하나가 ‘시간 놓치지 않기’입니다. "
        "운영 점검·이벤트 시각·결정전 같은 정해진 시간이 다가오면 Windows 토스트, "
        "인앱 팝업, 텔레그램으로 동시에 알려줍니다."
    )

    capture_placeholder(doc, 9,
                         "알람 탭 — Windows · Telegram · 인앱 토스트가 한 번에 묶인 화면",
                         image_name="05_alerts.png")

    h3(doc, "알람 하나에 들어가는 항목")
    bullet(doc, "시각", "시·분 단위 설정. 24시간 형식.")
    bullet(doc, "요일",
            "원하는 요일만 골라서 동작. 비워두면 매일 알림.")
    bullet(doc, "모드", "반복 — 매주 같은 요일에 반복 / 1회 — 가장 가까운 한 번만.")
    bullet(doc, "재알림",
            "알림이 떠도 못 보고 지나친 경우를 위해, N분 간격으로 다시 알려줍니다(0이면 한 번만).")
    bullet(doc, "라벨", "‘영지전 시작’, ‘점검 30분 전’ 등 식별용 이름.")

    h3(doc, "팝업 동작 설정 (전역)")
    bullet(doc, "재알림 시간",
            "재알림 기능을 켰을 때 몇 분 간격으로 다시 띄울지.")
    bullet(doc, "사운드",
            "기본(Windows 시스템음) / .wav 파일 지정 / 무음 중 선택.")
    bullet(doc, "볼륨", "0~100% (소리 직접 들으며 미리듣기 가능).")

    info_box(doc, "TIP", [
        "텔레그램과 함께 쓰면 자리를 비웠을 때도 휴대폰으로 알림이 옵니다.",
        "‘1회’ 모드는 일정 한 번 끝나면 자동으로 OFF 되니, 매번 새로 등록할 필요가 없습니다.",
    ])


# ───────── 섹션 10 키 입력 방식 ─────────
def section_10_input(doc):
    section_header(
        doc, "10", "키 입력 방식 선택",
        "환경에 맞춰 PostMessage 와 아두이노 HID 중에서 고릅니다.",
    )
    lead(
        doc,
        "매크로가 슬롯 키를 ‘어떻게’ 게임에 전달할지 정하는 부분입니다. "
        "각 방식은 동작 원리가 달라서 환경에 따라 잘 맞는 게 있고 그렇지 않은 게 있습니다."
    )

    h3(doc, "두 가지 방식 비교")
    compare_table(doc, rows=[
        ("방식", "PostMessage", "아두이노 HID"),
        ("원리",
         "Win32 PostMessage 로 게임 창에 직접 키 메시지를 전달",
         "USB 로 연결된 아두이노가 실제 HID 키보드처럼 키 입력 송신"),
        ("필수 장비", "없음 (소프트웨어만)", "아두이노 보드 (Pro Micro / Leonardo 등) + USB"),
        ("게임 창 포커스", "필요 없음 — 백그라운드 동작", "포커스 필요"),
        ("게임 측 인식", "윈도우 메시지 큐로 들어가는 소프트웨어 입력",
         "OS 가 인식하는 실제 키보드 입력"),
        ("도입 난이도", "설치 후 바로 사용", "보드 + 펌웨어 + COM 포트 설정 필요"),
        ("추천 상황",
         "PC 한 대로 가볍게 쓸 때",
         "장시간·안정 운영 우선 / 다중 환경"),
    ])

    h3(doc, "전환 방법")
    bullet(doc, "1단계", "설정 탭으로 이동.")
    bullet(doc, "2단계", "좌측 서브메뉴 ‘입력 방식’ 클릭.")
    bullet(doc, "3단계",
            "두 방식 중 하나 선택. 아두이노 선택 시 COM 포트와 보드레이트(기본 9600) 입력.")
    bullet(doc, "4단계",
            "‘테스트 키 전송’ 버튼으로 게임에 실제 입력이 들어가는지 확인.")

    info_box(doc, "TROUBLESHOOTING", [
        "PostMessage 가 게임에 안 닿는 것 같다면 → 게임 클라이언트의 보안 정책이 윈도우 메시지를 "
        "필터하는 경우입니다. 아두이노 방식으로 전환을 검토해보세요.",
        "아두이노가 인식되지 않는다면 → 장치 관리자에서 COM 포트 번호와 드라이버 설치 상태부터 "
        "확인합니다. CH340 계열 보드는 별도 드라이버가 필요합니다.",
    ])


# ───────── 섹션 11 텔레그램 · 편의 기능 ─────────
def section_11_extras(doc):
    section_header(
        doc, "11", "텔레그램 · 편의 기능",
        "원격 알림과 일상에서 작게 도움되는 도구들.",
    )
    lead(
        doc,
        "꼭 필요한 핵심 기능은 아니지만, 사냥을 좀 더 편하게 만들어주는 부가 기능들입니다. "
        "원격으로 상태를 받아보거나, UI를 정리하거나, 화면을 더 크게 보거나 — 필요할 때 켜서 쓰시면 됩니다."
    )

    capture_placeholder(doc, 10,
                         "설정 탭 — Arduino / Telegram 연결 + 캡처 FPS",
                         image_name="06_settings.png")

    h3(doc, "텔레그램 알림 연결")
    body(
        doc,
        "텔레그램 봇을 만들어 토큰을 등록하면, 매크로의 주요 이벤트(마스터 자동 OFF, "
        "사냥터 복귀 트리거, 특정 슬롯 발동 등)를 휴대폰으로 받아볼 수 있습니다."
    )
    bullet(doc, "봇 만들기",
            "텔레그램에서 @BotFather 로 새 봇 생성 → 발급된 API 토큰을 복사.")
    bullet(doc, "내 chat_id 확인",
            "텔레그램 @userinfobot 과 대화 시작 → 알려주는 숫자가 chat_id 입니다.")
    bullet(doc, "QuickCast 연결",
            "설정 → 연결 카드에서 토큰·chat_id 입력 후 ‘연결’ 클릭. 자동으로 테스트 메시지가 발송됩니다.")

    h3(doc, "ROI 잠금 (🔒)")
    body(
        doc,
        "사냥 중에 실수로 ROI 박스를 드래그해 좌표가 어긋나는 일을 막아주는 기능입니다. "
        "한 번 잘 맞춰 둔 박스는 잠가두는 게 안전합니다. 상단 자물쇠 아이콘으로 토글합니다."
    )

    h3(doc, "전체화면 미리보기 (F11)")
    body(
        doc,
        "라이브 캡처 화면을 큰 창으로 띄워 ROI 위치를 더 정밀하게 조정하고 싶을 때 유용합니다. "
        "F11 로 토글, ESC 로 빠져나옵니다."
    )

    h3(doc, "테마")
    body(
        doc,
        "다크 / 라이트 / 그라파이트 등 사용 환경에 맞춰 색감을 바꿀 수 있습니다. "
        "주야간 사용 시간이 다르면 가독성에 영향이 있으니 한 번 확인해 보세요. "
        "Ctrl+T 단축키로 즉시 전환됩니다."
    )

    info_box(doc, "ALSO", [
        "Ctrl+M — 마스터 ON/OFF 단축키.",
        "Ctrl+K — 명령 팔레트 열기 (메뉴 검색).",
        "Ctrl+Shift+T — 튜토리얼 다시 보기.",
    ])


# ───────── 섹션 12 FAQ ─────────
def section_12_faq(doc):
    section_header(
        doc, "12", "자주 묻는 질문",
        "안 될 때 가장 먼저 확인하는 곳.",
    )

    h3(doc, "Q. 매크로를 켰는데 아무 키도 안 나가요.")
    body(
        doc,
        "1) 마스터 스위치가 ON 인지 확인. 2) 슬롯이 활성(스위치 빨강)인지 확인. "
        "3) HP·MP 라이브 수치가 캐릭터 실제 값과 일치하는지 확인 — 일치하지 않으면 캡처 ROI 가 "
        "어긋나 슬롯 발동 조건을 충족하지 못합니다. 4) 설정 → 입력 방식에서 ‘테스트 키 전송’ 으로 "
        "키 입력 경로 자체가 게임에 닿는지 확인."
    )

    h3(doc, "Q. PK 표시가 잘 안 잡혀요.")
    body(
        doc,
        "PK ROI 박스가 캐릭터 머리 위 영역을 충분히 덮고 있는지 확인합니다. 캐릭터 위치가 "
        "조금씩 움직이므로 박스를 ‘검색 영역 모드’ 로 살짝 크게 잡는 것이 안정적입니다. "
        "감지 점수가 임계값 근처에서 흔들리면 임계값을 +100,000 정도 낮춰 보세요."
    )

    h3(doc, "Q. 펫 호루라기가 가끔 엉뚱한 순간에 ESC 가 나갑니다.")
    body(
        doc,
        "v1.0.3 에서는 3초 유지 트리거가 적용되어 짧은 오탐은 차단되지만, ROI 가 너무 넓거나 "
        "민감도가 낮으면 여전히 오탐 가능. 발바닥 ROI 박스를 발바닥 위치에만 타이트하게 잡고, "
        "감지 민감도를 70% 이상으로 올려보세요."
    )

    h3(doc, "Q. 사냥터 복귀 시퀀스가 빈 곳을 클릭해요.")
    body(
        doc,
        "보통 단계 사이 ‘대기 시간’ 이 짧아 화면 전환이 끝나기 전에 클릭이 발사된 경우입니다. "
        "각 단계의 대기 시간을 1000~3000ms 정도 넉넉히 늘려 보세요. 마을 도착 직후의 안전지대 "
        "버프가 풀리지 않아서일 수도 있는데, 그건 ‘시작 지연(start delay)’ 값을 늘리면 됩니다."
    )

    h3(doc, "Q. 다른 PC 로 옮길 때 설정을 그대로 옮길 수 있나요?")
    body(
        doc,
        "userdata.json 파일 하나를 복사해 새 PC 의 같은 위치에 두면 그대로 복원됩니다. "
        "설정 → 데이터 카드에서 백업·복원 버튼으로 파일 위치를 열어볼 수도 있습니다."
    )

    h3(doc, "Q. 안티치트나 약관에 문제가 되지 않나요?")
    body(
        doc,
        "리니지W의 약관은 자동 입력 도구 사용을 일반적으로 제한합니다. 본 글은 학습·참고용 "
        "안내이며, 실사용 여부와 그에 따른 책임은 전적으로 사용자 본인에게 있습니다. "
        "안티치트 정책은 시기마다 다를 수 있으니 운영 정책을 직접 확인하시기 바랍니다."
    )


# ───────── 백 커버 ─────────
def make_back_cover(doc):
    add_page_break(doc)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(17.0)
    cell = tbl.cell(0, 0)
    cell.width = Cm(17.0)
    set_cell_shading(cell, "1A1A1A")

    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, mm in (("top", 2800), ("left", 700), ("bottom", 2800), ("right", 700)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(mm))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(8)
    r0 = p0.add_run("DOWNLOAD  &  CONTACT")
    style_run(r0, font=FONT_EN, size_pt=10, bold=True,
              color=GOLD, letter_spacing_pt=2.5)

    p1 = cell.add_paragraph()
    p1.paragraph_format.space_before = Pt(8)
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run("받아보기")
    style_run(r1, size_pt=32, bold=True,
              color=RGBColor(0xFF, 0xFF, 0xFF))

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(20)
    r2 = p2.add_run("━━━━━")
    style_run(r2, font=FONT_EN, size_pt=14, bold=True, color=ACCENT)

    # 다운로드 항목 (링크 포함)
    p3 = cell.add_paragraph()
    p3.paragraph_format.space_after = Pt(6)
    r3a = p3.add_run("·  ")
    style_run(r3a, size_pt=11, color=RGBColor(0xDD, 0xDD, 0xDD))
    r3b = p3.add_run("최신 릴리즈 ")
    style_run(r3b, size_pt=11, color=RGBColor(0xDD, 0xDD, 0xDD), bold=True)
    add_hyperlink(
        p3,
        "https://github.com/i0500/quickcast/releases/tag/v1.0.3",
        "github.com/i0500/quickcast/releases/tag/v1.0.3",
        size_pt=11,
        color=GOLD,
    )

    p4 = cell.add_paragraph()
    p4.paragraph_format.space_after = Pt(6)
    r4a = p4.add_run("·  ")
    style_run(r4a, size_pt=11, color=RGBColor(0xDD, 0xDD, 0xDD))
    r4b = p4.add_run("바로 다운로드 ")
    style_run(r4b, size_pt=11, color=RGBColor(0xDD, 0xDD, 0xDD), bold=True)
    add_hyperlink(
        p4,
        "https://github.com/i0500/quickcast/releases/download/v1.0.3/quickcast-v1.0.3.zip",
        "quickcast-v1.0.3.zip",
        size_pt=11,
        color=GOLD,
    )
    r4c = p4.add_run("  (약 290MB · 비밀번호 0000)")
    style_run(r4c, size_pt=10.5, color=RGBColor(0xAA, 0xAA, 0xAA))

    p5 = cell.add_paragraph()
    p5.paragraph_format.space_before = Pt(28)
    p5.paragraph_format.space_after = Pt(4)
    r5 = p5.add_run("긴 글 읽어주셔서 감사합니다.")
    style_run(r5, size_pt=13, color=RGBColor(0xFF, 0xFF, 0xFF),
              letter_spacing_pt=0.5)

    p6 = cell.add_paragraph()
    p6.paragraph_format.space_after = Pt(28)
    r6 = p6.add_run("질문 / 제안 / 버그 제보는 GitHub Issues 또는 댓글로 부탁드립니다.")
    style_run(r6, size_pt=11, color=RGBColor(0xCC, 0xCC, 0xCC))

    p7 = cell.add_paragraph()
    r7 = p7.add_run("#리니지W  #리니지W매크로  #QuickCast  #자동매크로  #리니지W자동사냥")
    style_run(r7, size_pt=10, color=RGBColor(0x99, 0x99, 0x99),
              letter_spacing_pt=0.3)


# ───────── 빌드 ─────────
def main():
    doc = Document()
    setup_page(doc)

    # 1 — 표지 + 도입부
    make_cover(doc)
    make_intro(doc)

    # 2 — 섹션 01 ~ 04
    section_01_compare(doc)
    section_02_install(doc)
    section_03_power(doc)
    section_04_hpmp(doc)

    # 3 — 섹션 05 ~ 08
    section_05_slots(doc)
    section_06_pk_potion(doc)
    section_07_pet_whistle(doc)
    section_08_recovery(doc)

    # 4 — 섹션 09 ~ 12
    section_09_alerts(doc)
    section_10_input(doc)
    section_11_extras(doc)
    section_12_faq(doc)

    # 백 커버
    make_back_cover(doc)

    out = Path(r"F:/린w/dist/QuickCast_v1.0.3_네이버포스트.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(out))
        print(f"saved: {out}  size={out.stat().st_size}")
    except PermissionError:
        alt = out.with_name(out.stem + "_new.docx")
        doc.save(str(alt))
        print(f"원본이 Word에 열려있어 임시 저장: {alt}")


if __name__ == "__main__":
    main()
