"""
QuickCast 네이버 포스트용 DOCX 생성 — Step 1 (독자 친화 톤)
표지 + 도입부 + 목차.
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ───────── 디자인 토큰 ─────────
FONT_KR = "맑은 고딕"
FONT_EN = "Segoe UI"

INK = RGBColor(0x1A, 0x1A, 0x1A)
INK_SOFT = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0xB3, 0x2C, 0x2C)         # 와인 레드
RULE = RGBColor(0xD9, 0xD2, 0xC4)
GOLD = RGBColor(0xC8, 0x96, 0x5A)


# ───────── XML 유틸 ─────────
def set_cell_shading(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_page_break(doc) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._r.append(br)


def style_run(run, *, font=FONT_KR, size_pt=11.0, bold=False, italic=False,
              color=INK, letter_spacing_pt=None):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_KR)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    if letter_spacing_pt is not None:
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), str(int(letter_spacing_pt * 20)))
        rPr.append(spacing)


def add_paragraph(doc, text="", *, align=None, space_before=0, space_after=4,
                  line_spacing=1.6, **run_kwargs):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        style_run(r, **run_kwargs)
    return p


# ───────── 표지 ─────────
def make_cover(doc):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(17.0)
    cell = tbl.cell(0, 0)
    cell.width = Cm(17.0)
    set_cell_shading(cell, "1A1A1A")

    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, mm in (("top", 3400), ("left", 700), ("bottom", 3400), ("right", 700)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(mm))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(8)
    r0 = p0.add_run("GUIDE  v1.0.3  ·  MAY 2026")
    style_run(r0, font=FONT_EN, size_pt=10, color=GOLD,
              letter_spacing_pt=2.5, bold=True)

    p1 = cell.add_paragraph()
    p1.paragraph_format.space_before = Pt(8)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run("QuickCast")
    style_run(r1, font=FONT_EN, size_pt=64, bold=True,
              color=RGBColor(0xFF, 0xFF, 0xFF), letter_spacing_pt=-1.0)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(28)
    r2 = p2.add_run("리니지W 화면 인식 매크로 사용설명서")
    style_run(r2, size_pt=21, color=RGBColor(0xFF, 0xFF, 0xFF),
              letter_spacing_pt=0.4)

    p3 = cell.add_paragraph()
    p3.paragraph_format.space_after = Pt(20)
    r3 = p3.add_run("━━━━━")
    style_run(r3, font=FONT_EN, size_pt=14, bold=True, color=ACCENT)

    p4 = cell.add_paragraph()
    r4 = p4.add_run("HP·MP 상황별 스킬 자동 사용  ·  PK 대응  ·  사냥터 복귀")
    style_run(r4, size_pt=12, color=RGBColor(0xDD, 0xDD, 0xDD),
              letter_spacing_pt=0.3)
    p5 = cell.add_paragraph()
    p5.paragraph_format.space_before = Pt(2)
    r5 = p5.add_run("아두이노 + 웹 기반 매크로의 데스크톱 버전")
    style_run(r5, size_pt=12, color=RGBColor(0xDD, 0xDD, 0xDD),
              letter_spacing_pt=0.3)

    p6 = cell.add_paragraph()
    p6.paragraph_format.space_before = Pt(36)
    r6 = p6.add_run("github.com/i0500/quickcast")
    style_run(r6, font=FONT_EN, size_pt=10, color=RGBColor(0x99, 0x99, 0x99),
              letter_spacing_pt=1.2)


# ───────── 도입부 ─────────
def make_intro(doc):
    add_page_break(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("들어가며")
    style_run(r, font=FONT_EN, size_pt=10, bold=True,
              color=ACCENT, letter_spacing_pt=2.5)

    add_paragraph(doc, "QuickCast란?",
                  size_pt=24, bold=True, line_spacing=1.2, space_after=14)

    add_paragraph(
        doc,
        "QuickCast는 게임 화면을 실시간으로 인식해 PostMessage 방식으로 키 입력을 보내는 "
        "리니지W 보조 매크로입니다. 기존 아두이노 + 웹 브라우저 기반 매크로의 동작 원리를 "
        "유지하면서, 데스크톱 단독 실행이 가능한 형태로 다시 구현한 버전입니다.",
        size_pt=11.5, line_spacing=1.75, space_after=10,
    )
    add_paragraph(
        doc,
        "물약과 버프는 리니지W가 자체적으로 자동 사용하기 때문에 매크로가 다루는 영역이 아닙니다. "
        "QuickCast가 담당하는 것은 그 위의 ‘상황 대응’입니다. HP·MP가 특정 구간일 때만 "
        "필요한 스킬을 자동으로 사용하고, PK 감지 시 정해둔 키를 입력하며, 물약이 떨어졌을 때는 "
        "사냥터 복귀 시퀀스를 자동으로 실행합니다.",
        size_pt=11.5, line_spacing=1.75, space_after=10,
    )
    add_paragraph(
        doc,
        "동작 방식은 단순합니다. 게임 창을 주기적으로 캡처해 체력 바·마나 바·PK 표시·물약 부족(!) 같은 "
        "요소를 OpenCV 템플릿 매칭으로 인식하고, 설정한 조건이 충족되면 PostMessage(또는 아두이노 HID) "
        "로 키 입력을 보냅니다. 게임 창에 포커스가 없어도 동작하며, 입력 방식은 환경에 따라 선택할 수 있습니다.",
        size_pt=11.5, line_spacing=1.75, space_after=10,
    )
    add_paragraph(
        doc,
        "이 글은 v1.0.3 기준 전체 기능을 항목별로 정리한 사용설명서입니다. "
        "다음 장에서 기존 웹 버전과 어떤 점이 달라졌는지 비교한 뒤, 각 기능을 차례로 살펴봅니다.",
        size_pt=11.5, line_spacing=1.75, space_after=22,
    )

    add_paragraph(doc, "─" * 38, color=RULE, size_pt=10, space_after=18)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("목차")
    style_run(r, font=FONT_EN, size_pt=10, bold=True,
              color=ACCENT, letter_spacing_pt=2.5)

    toc = [
        ("01", "기존 웹 버전과 다른 점",      "동작 원리는 같고, 환경이 달라졌습니다"),
        ("02", "다운로드 · 설치",            "받는 곳, 압축 비밀번호, 첫 실행"),
        ("03", "매크로 켜기 · 끄기",          "마스터 스위치와 플로터"),
        ("04", "HP · MP 인식 설정",          "캡처 박스 위치, OCR 텍스트 모드"),
        ("05", "슬롯 — 상황별 스킬 자동 사용", "HP·MP 구간에 맞춰 키 입력"),
        ("06", "PK · 물약 부족 대응",        "위급 상황 자동 응답"),
        ("07", "펫 호루라기 자동 닫기 NEW",   "v1.0.3 — 3초 유지 트리거"),
        ("08", "사냥터 자동 복귀",           "마을 → 사냥터 시퀀스"),
        ("09", "알람",                       "시간 기반 알림"),
        ("10", "키 입력 방식 선택",           "PostMessage 또는 아두이노 HID"),
        ("11", "텔레그램 · 편의 기능",        "원격 알림, ROI 잠금, 전체화면"),
        ("12", "자주 묻는 질문",             "안 될 때 확인하는 곳"),
    ]
    for num, name, desc in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.35
        r1 = p.add_run(num + "  ")
        style_run(r1, font=FONT_EN, size_pt=11, bold=True, color=ACCENT)
        r2 = p.add_run(name)
        style_run(r2, size_pt=12, bold=True, color=INK)
        r3 = p.add_run("   " + desc)
        style_run(r3, size_pt=10.5, color=INK_SOFT)


def setup_page(doc):
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    style = doc.styles["Normal"]
    style.font.name = FONT_KR
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_KR)
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)


def main():
    doc = Document()
    setup_page(doc)
    make_cover(doc)
    make_intro(doc)

    out = Path(r"F:/린w/dist/QuickCast_v1.0.3_네이버포스트.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(out))
        print(f"saved: {out}  size={out.stat().st_size}")
    except PermissionError:
        alt = out.with_name(out.stem + "_new.docx")
        doc.save(str(alt))
        print(f"원본이 Word에 열려있어 임시 저장: {alt}")


if __name__ == "__main__":
    main()
