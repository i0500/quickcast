"""
QuickCast 네이버 포스트용 DOCX — Step 2
섹션 01 기존 웹 버전과 다른 점
섹션 02 다운로드 · 설치
섹션 03 매크로 켜기 · 끄기
섹션 04 HP · MP 인식 설정
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Step 1 의 헬퍼/디자인 토큰을 그대로 재사용
from make_naverpost_step1 import (
    FONT_KR, FONT_EN, INK, INK_SOFT, ACCENT, RULE, GOLD,
    set_cell_shading, add_page_break, style_run, add_paragraph,
    make_cover, make_intro, setup_page,
)


# ───────── 추가 디자인 헬퍼 ─────────
def section_header(doc, num: str, title: str, subtitle: str = "") -> None:
    """매거진 섹션 헤더 — 큰 회색 숫자 / 와인레드 막대 / 제목 / 부제"""
    add_page_break(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(num)
    style_run(r, font=FONT_EN, size_pt=42, bold=True,
              color=RGBColor(0xE0, 0xD8, 0xC8), letter_spacing_pt=-1.0)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("━━")
    style_run(r, font=FONT_EN, size_pt=12, bold=True, color=ACCENT)

    add_paragraph(doc, title, size_pt=22, bold=True,
                  line_spacing=1.2, space_after=4)
    if subtitle:
        add_paragraph(doc, subtitle, size_pt=11.5, color=INK_SOFT,
                      line_spacing=1.4, space_after=18)


def lead(doc, text: str) -> None:
    """섹션 도입 한 단락 — 살짝 큰 글씨, 짙은 회색"""
    add_paragraph(doc, text, size_pt=12.5, color=RGBColor(0x33, 0x33, 0x33),
                  line_spacing=1.7, space_after=12)


def body(doc, text: str) -> None:
    add_paragraph(doc, text, size_pt=11, color=INK,
                  line_spacing=1.75, space_after=8)


def small(doc, text: str) -> None:
    add_paragraph(doc, text, size_pt=10, color=INK_SOFT,
                  line_spacing=1.5, space_after=8)


def h3(doc, text: str) -> None:
    """소제목 — 작은 빨간 점 + 텍스트"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run("● ")
    style_run(r1, font=FONT_EN, size_pt=8, bold=True, color=ACCENT)
    r2 = p.add_run(text)
    style_run(r2, size_pt=13, bold=True, color=INK, letter_spacing_pt=0.2)


def bullet(doc, label: str, text: str) -> None:
    """라벨이 굵은 인라인 불릿"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.65
    p.paragraph_format.left_indent = Cm(0.4)
    r1 = p.add_run("·  ")
    style_run(r1, size_pt=11, color=INK_SOFT, bold=True)
    r2 = p.add_run(label)
    style_run(r2, size_pt=11, color=INK, bold=True)
    r3 = p.add_run("  " + text)
    style_run(r3, size_pt=11, color=INK)


CAPTURES_DIR = Path(r"F:/린w/dist/captures")


# ───────── 하이퍼링크 ─────────
def add_hyperlink(paragraph, url: str, text: str, *,
                   size_pt: float = 11.0, bold: bool = False,
                   color: RGBColor = ACCENT) -> None:
    """Word 본문에 클릭 가능한 외부 링크를 추가한다."""
    from docx.opc.constants import RELATIONSHIP_TYPE
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), FONT_KR)
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)

    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)

    c = OxmlElement("w:color")
    c.set(qn("w:val"), str(color))   # RGBColor → 'B32C2C' 6자리 hex
    rPr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def link_bullet(doc, label: str, link_text: str, url: str,
                trailing: str = "") -> None:
    """굵은 라벨 + 클릭 가능한 링크 + 보조 텍스트."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.65
    p.paragraph_format.left_indent = Cm(0.4)
    r1 = p.add_run("·  ")
    style_run(r1, size_pt=11, color=INK_SOFT, bold=True)
    r2 = p.add_run(label)
    style_run(r2, size_pt=11, color=INK, bold=True)
    r3 = p.add_run("  ")
    style_run(r3, size_pt=11, color=INK)
    add_hyperlink(p, url, link_text, size_pt=11, color=ACCENT)
    if trailing:
        r4 = p.add_run("  " + trailing)
        style_run(r4, size_pt=11, color=INK_SOFT)


def capture_placeholder(doc, num: int, caption: str, image_name: str = "") -> None:
    """캡처 자리.
    image_name(예: '02_capture.png') 으로 captures 폴더의 파일을 찾아 본문에 삽입.
    파일이 없으면 크림색 placeholder 박스로 대체.
    """
    img_path = CAPTURES_DIR / image_name if image_name else None
    if img_path and img_path.exists():
        # 가운데 정렬된 이미지 + 캡션
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run()
        r.add_picture(str(img_path), width=Cm(15.5))

        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_after = Pt(14)
        cr1 = cap_p.add_run(f"📷 {num:02d}  ")
        style_run(cr1, font=FONT_EN, size_pt=9.5, bold=True, color=ACCENT,
                  letter_spacing_pt=1.5)
        cr2 = cap_p.add_run(caption)
        style_run(cr2, size_pt=10, color=INK_SOFT, italic=True)
        return

    # 파일 없음 → placeholder
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(15.5)
    cell = tbl.cell(0, 0)
    cell.width = Cm(15.5)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(cell, "F5F1EA")

    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, mm in (("top", 1600), ("left", 400), ("bottom", 1600), ("right", 400)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(mm))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"📷  CAPTURE  {num:02d}")
    style_run(r, font=FONT_EN, size_pt=11, bold=True, color=ACCENT,
              letter_spacing_pt=2.5)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(6)
    r2 = p2.add_run(caption)
    style_run(r2, size_pt=11, color=INK_SOFT)

    add_paragraph(doc, "", space_after=12)


def info_box(doc, title: str, lines: list[str]) -> None:
    """크림색 사이드 박스 — 좌측 와인레드 막대 느낌은 텍스트 색상으로 표현"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(15.5)
    cell = tbl.cell(0, 0)
    cell.width = Cm(15.5)
    set_cell_shading(cell, "F5F1EA")

    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, mm in (("top", 350), ("left", 500), ("bottom", 350), ("right", 500)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(mm))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    style_run(r, font=FONT_EN, size_pt=10, bold=True, color=ACCENT,
              letter_spacing_pt=2.0)

    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.6
        r = p.add_run(line)
        style_run(r, size_pt=10.5, color=INK)

    add_paragraph(doc, "", space_after=10)


def compare_table(doc, rows: list[tuple[str, str, str]]) -> None:
    """3열 비교표 — 헤더: 항목 / 기존 웹 버전 / QuickCast"""
    tbl = doc.add_table(rows=len(rows) + 1, cols=3)
    tbl.autofit = False
    widths = (Cm(4.0), Cm(5.6), Cm(5.9))
    for i, w in enumerate(widths):
        for cell in tbl.columns[i].cells:
            cell.width = w

    def _cell_borders(cell, fill="auto"):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "bottom"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "6")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "1A1A1A")
            tcBorders.append(b)
        for side in ("left", "right"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "nil")
            tcBorders.append(b)
        tcPr.append(tcBorders)

    # 헤더
    hdr_cells = tbl.rows[0].cells
    headers = ("항목", "기존 웹 버전", "QuickCast")
    for c, txt in zip(hdr_cells, headers):
        set_cell_shading(c, "1A1A1A")
        _cell_borders(c)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(txt)
        style_run(r, size_pt=10.5, bold=True,
                  color=RGBColor(0xFF, 0xFF, 0xFF), letter_spacing_pt=1.0)

    # 데이터
    for i, (col1, col2, col3) in enumerate(rows, start=1):
        row_cells = tbl.rows[i].cells
        for c, txt, bold in ((row_cells[0], col1, True),
                              (row_cells[1], col2, False),
                              (row_cells[2], col3, False)):
            _cell_borders(c)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.45
            r = p.add_run(txt)
            style_run(r, size_pt=10.5, color=INK, bold=bold)

    add_paragraph(doc, "", space_after=8)


# ───────── 섹션 본문 ─────────
def section_01_compare(doc):
    section_header(
        doc, "01", "기존 웹 버전과 다른 점",
        "동작 원리는 같고, 환경이 달라졌습니다.",
    )
    lead(
        doc,
        "QuickCast는 새로운 매크로가 아니라, 기존 아두이노 + 웹 브라우저 기반 매크로의 "
        "데스크톱 재구현 버전입니다. 화면을 보고 키를 보낸다는 동작 원리는 그대로 두고, "
        "실행 환경과 입력 방식의 선택지를 넓힌 형태입니다."
    )

    h3(doc, "한눈에 비교")
    compare_table(doc, rows=[
        ("실행 형태", "브라우저(웹페이지)에서 실행",
         "단독 실행 .exe (브라우저 불필요)"),
        ("필수 장비", "아두이노 보드 + USB 연결",
         "선택 — 아두이노 없이도 동작"),
        ("입력 방식", "아두이노 HID 전용",
         "PostMessage / 아두이노 (2종 선택)"),
        ("게임 창 포커스", "포커스 필요한 경우 많음",
         "PostMessage 모드에서는 백그라운드 동작"),
        ("화면 인식 엔진", "JavaScript + OpenCV.js",
         "Python + OpenCV 네이티브 (체감 속도 개선)"),
        ("게임 창 자동 인식", "수동 지정",
         "5초 주기 자동 탐색·재연결"),
        ("ROI 시각 편집", "수치 입력 위주",
         "라이브 프리뷰에서 드래그·리사이즈"),
        ("설정 보관", "브라우저 로컬스토리지",
         "userdata.json 파일 (백업·이식 용이)"),
    ])

    body(
        doc,
        "기능 자체는 거의 동일하게 옮겨왔습니다. 기존 사용자가 보던 슬롯·HP·MP·PK·사냥터 복귀 같은 "
        "개념이 그대로 남아 있고, 거기에 데스크톱 환경에서 더 잘 동작하도록 손본 부분이 "
        "더해진 정도로 보시면 됩니다."
    )

    info_box(doc, "NOTE", [
        "아두이노가 있다면 그대로 사용할 수 있습니다. 백엔드만 ‘arduino’로 선택하면 됩니다.",
        "아두이노가 없거나 게임 정책상 사용이 어렵다면 PostMessage 또는 AttachInput으로 전환 가능.",
    ])


def section_02_install(doc):
    section_header(
        doc, "02", "다운로드 · 설치",
        "별도 설치 과정 없이 압축만 풀면 됩니다.",
    )
    lead(
        doc,
        "QuickCast는 단일 실행 파일 형태로 배포됩니다. 설치 프로그램은 따로 없고, "
        "내려받은 zip의 압축을 풀면 그 안의 quickcast.exe를 더블클릭하는 것으로 끝납니다."
    )

    h3(doc, "받는 곳")
    link_bullet(
        doc, "릴리즈 페이지",
        "github.com/i0500/quickcast/releases/tag/v1.0.3",
        "https://github.com/i0500/quickcast/releases/tag/v1.0.3",
        trailing="— 최신 버전 v1.0.3",
    )
    link_bullet(
        doc, "바로 다운로드",
        "quickcast-v1.0.3.zip (약 290MB)",
        "https://github.com/i0500/quickcast/releases/download/v1.0.3/quickcast-v1.0.3.zip",
    )
    bullet(doc, "압축 비밀번호", "0000")
    bullet(doc, "권장 도구", "7-Zip / 반디집 / WinRAR")

    capture_placeholder(doc, 1,
                         "GitHub 릴리즈 페이지 — 상단의 다운로드 버튼이 보이는 영역")

    h3(doc, "권장 환경")
    bullet(doc, "운영체제", "Windows 10 / 11")
    bullet(doc, "리니지W 해상도", "최저 1280×720 이상 권장. 더 높은 해상도도 동작하며, 해상도가 바뀌면 ROI만 한 번 다시 맞추면 됩니다.")
    bullet(doc, "디스플레이 배율", "100% 권장. 125·150% 환경에서도 동작하나 캡처 좌표 보정이 필요할 수 있습니다.")

    h3(doc, "첫 실행 흐름")
    body(
        doc,
        "압축을 풀고 quickcast.exe를 실행하면 메인 창이 열립니다. 처음 실행 시 화면 위에 "
        "튜토리얼 안내가 단계별로 표시되며, 건너뛰거나 따라가며 기본 사용법을 익힐 수 있습니다. "
        "튜토리얼은 도움말 메뉴에서 다시 열 수 있습니다."
    )

    capture_placeholder(doc, 2,
                         "QuickCast 메인 창 — 대시보드 탭이 열린 전체 화면",
                         image_name="01_dashboard.png")

    info_box(doc, "TIP", [
        "처음에는 게임 창과 QuickCast 창을 모니터에 나란히 띄워두면 ROI 조정이 편합니다.",
        "userdata.json 파일에 모든 설정이 저장됩니다. 다른 PC로 옮길 때 이 파일만 복사하면 됩니다.",
    ])


def section_03_power(doc):
    section_header(
        doc, "03", "매크로 켜기 · 끄기",
        "마스터 스위치와 게임 창 옆 플로터.",
    )
    lead(
        doc,
        "QuickCast 전체 동작은 상단의 ‘마스터 스위치’ 하나로 ON·OFF 됩니다. "
        "켜진 동안 캡처 루프와 슬롯 로직이 함께 동작하고, 끄면 즉시 멈춥니다."
    )

    h3(doc, "마스터 스위치")
    bullet(doc, "ON",
            "3초 카운트다운 후 모든 감지·슬롯 동작 시작. "
            "잘못 켰을 때 빠져나올 시간을 주기 위한 안전 그레이스 타임.")
    bullet(doc, "OFF",
            "즉시 모든 동작 정지. 진행 중이던 사냥터 복귀 시퀀스도 함께 중단.")
    bullet(doc, "재시작 시", "안전을 위해 항상 OFF 상태로 시작.")

    capture_placeholder(doc, 3,
                         "상단 우측의 Floating · Master 토글 — ON 상태",
                         image_name="08_toggles.png")

    h3(doc, "게임 창 옆 플로터")
    body(
        doc,
        "마스터 스위치는 메인 창에 있지만, 사냥 중에는 게임 창에 시선이 가 있기 마련입니다. "
        "그래서 작은 위젯(플로터)이 게임 창 모서리에 떠 있어 한 번의 클릭으로 매크로 전체를 "
        "켜고 끌 수 있습니다. ▼ 버튼을 누르면 패널이 펼쳐져 슬롯별 토글도 직접 조작할 수 있습니다."
    )

    capture_placeholder(doc, 4,
                         "게임 창 우측에 붙은 플로터 — 닫힌 상태와 펼친 상태")

    info_box(doc, "참고", [
        "플로터는 항상 표시 / 자동 부착 옵션이 있어 게임 창을 따라다닙니다.",
        "게임 창이 최소화되거나 사라지면 자동으로 숨겨졌다가, 다시 발견되면 재부착됩니다.",
    ])


def section_04_hpmp(doc):
    section_header(
        doc, "04", "HP · MP 인식 설정",
        "캐릭터 체력·마나를 정확하게 읽도록 박스를 맞춥니다.",
    )
    lead(
        doc,
        "QuickCast의 모든 자동 동작은 ‘지금 HP가 몇 퍼센트인가, MP가 몇 퍼센트인가’를 "
        "정확히 읽는 데서 출발합니다. 캡처 영역(ROI)이 체력 바 위에 잘 맞아 있어야 슬롯이 의도대로 동작합니다."
    )

    h3(doc, "기본 위치")
    body(
        doc,
        "1280×720 창 모드 기준으로 좌측 상단의 체력 바·마나 바 위치에 자동 정렬된 기본값이 들어 있습니다. "
        "이 해상도를 그대로 쓰신다면 별도 조정 없이 바로 사용 가능합니다."
    )

    h3(doc, "다른 해상도 / UI 스케일 사용 시")
    bullet(doc, "1단계",
            "캡처 탭으로 이동. 라이브 프리뷰 화면 위에 HP·MP 박스(분홍·파랑)가 표시됩니다.")
    bullet(doc, "2단계",
            "박스를 드래그해 체력 바·마나 바 위로 옮기고, 모서리를 잡아 폭·높이 조정.")
    bullet(doc, "3단계",
            "박스가 빨간(체력)·파란(마나) 픽셀만 정확히 덮도록 위·아래로 좁게 조정.")
    bullet(doc, "4단계",
            "상단 라이브 수치가 캐릭터 실제 HP·MP %와 일치하는지 확인.")

    capture_placeholder(doc, 5,
                         "캡처 탭 — HP·MP ROI 박스 등 캡처 관련 설정이 모인 화면",
                         image_name="02_capture.png")

    h3(doc, "수라 모드")
    body(
        doc,
        "수라(이프리트) 캐릭터는 체력 바가 미세하게 아래로 밀려 있습니다. "
        "설정에서 ‘수라 모드’를 켜면 캡처 좌표가 자동으로 그만큼 보정되어, ROI를 다시 그릴 필요가 없습니다."
    )

    h3(doc, "OCR 텍스트 모드 (고급)")
    body(
        doc,
        "체력 바가 가려지거나 색상이 변하는 던전·이펙트 환경에서는, 체력 바 대신 "
        "HP·MP·물약 ‘숫자’를 직접 읽는 OCR 모드를 사용할 수 있습니다."
    )
    bullet(doc, "학습", "캡처 탭의 OCR 카드를 펼친 뒤 각 숫자(0~9, %) 영역을 한 번씩 지정해 학습.")
    bullet(doc, "활성화", "학습이 끝나면 OCR 모드를 ON. 다음 프레임부터 OCR 결과가 우선 적용.")
    bullet(doc, "복귀", "OFF 하면 기존 체력 바 인식 방식으로 자동 복귀.")

    info_box(doc, "CHECKLIST", [
        "체력 바 모양이 바뀌는 버프(예: 보호막)에서도 % 가 정상 표시되는지 확인.",
        "박스 폭은 약간 좁게 — 다른 색이 박스에 섞이면 오인식 위험이 있습니다.",
        "라이브 수치가 흔들리면 박스 위치보다 ‘박스 높이’를 한 픽셀 줄여보세요.",
    ])


# ───────── 빌드 ─────────
def main():
    doc = Document()
    setup_page(doc)

    # Step 1 — 표지 + 도입부
    make_cover(doc)
    make_intro(doc)

    # Step 2 — 섹션 01 ~ 04
    section_01_compare(doc)
    section_02_install(doc)
    section_03_power(doc)
    section_04_hpmp(doc)

    out = Path(r"F:/린w/dist/QuickCast_v1.0.3_네이버포스트.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(out))
        print(f"saved: {out}  size={out.stat().st_size}")
    except PermissionError:
        alt = out.with_name(out.stem + "_new.docx")
        doc.save(str(alt))
        print(f"원본이 Word에 열려있어 임시 저장: {alt}")


if __name__ == "__main__":
    main()
